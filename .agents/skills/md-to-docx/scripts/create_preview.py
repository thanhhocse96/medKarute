#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(os.path.dirname(SCRIPT_DIR), 'assets')
TEMPLATE_PATH = os.path.join(ASSETS_DIR, 'template.docx')
OUTPUT_PATH = os.path.join(ASSETS_DIR, 'template_preview.docx')
CJK_FONT = 'SimSun'

def set_run_font(run, font_cn=CJK_FONT, font_en='Times New Roman', size=12, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)

def add_shading(paragraph, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    paragraph._p.get_or_add_pPr().append(shading)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_preview():
    doc = Document(TEMPLATE_PATH)

    doc.add_heading('Formatting Specification Preview', level=1)
    p = doc.add_paragraph()
    run = p.add_run('This document demonstrates every format supported by the md-to-docx skill.')
    set_run_font(run)

    doc.add_heading('1. Heading Styles', level=1)

    doc.add_heading('Level 1 Heading (Heading 1)', level=1)
    p = doc.add_paragraph()
    run = p.add_run('Size: 22pt, bold, page break before each level-1 heading')
    set_run_font(run, size=10, italic=True)

    doc.add_heading('Level 2 Heading (Heading 2)', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Size: 16pt, bold, no page break')
    set_run_font(run, size=10, italic=True)

    doc.add_heading('Level 3 Heading (Heading 3)', level=3)
    p = doc.add_paragraph()
    run = p.add_run('Size: 15pt, bold, no page break')
    set_run_font(run, size=10, italic=True)

    doc.add_heading('Level 4 Heading (Heading 4)', level=4)
    p = doc.add_paragraph()
    run = p.add_run('Size: 14pt, bold, no page break')
    set_run_font(run, size=10, italic=True)

    doc.add_heading('Level 5 Heading (Heading 5)', level=5)
    p = doc.add_paragraph()
    run = p.add_run('Size: 14pt, bold, no page break')
    set_run_font(run, size=10, italic=True)

    doc.add_heading('Level 6 Heading (Heading 6)', level=6)
    p = doc.add_paragraph()
    run = p.add_run('Size: 12pt, bold, no page break')
    set_run_font(run, size=10, italic=True)

    doc.add_heading('2. Body Text Styles', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Body: SimSun + Times New Roman, 12pt, first-line indent 0.74cm, 1.5× line spacing.')
    set_run_font(run)

    p = doc.add_paragraph()
    run = p.add_run('Bold text: ')
    set_run_font(run)
    run = p.add_run('This is bold text')
    set_run_font(run, bold=True)

    p = doc.add_paragraph()
    run = p.add_run('Italic text: ')
    set_run_font(run)
    run = p.add_run('This is italic text')
    set_run_font(run, italic=True)

    p = doc.add_paragraph()
    run = p.add_run('Bold + italic: ')
    set_run_font(run)
    run = p.add_run('This is bold italic text')
    set_run_font(run, bold=True, italic=True)

    doc.add_heading('3. List Styles', level=1)

    doc.add_heading('3.1 Unordered List', level=2)
    for item in ['Unordered item 1', 'Unordered item 2', 'Unordered item 3']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        set_run_font(run)
        p.paragraph_format.left_indent = Cm(0.74)

    doc.add_heading('3.2 Ordered List', level=2)
    for i, item in enumerate(['Ordered item 1', 'Ordered item 2', 'Ordered item 3'], 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {item}')
        set_run_font(run)
        p.paragraph_format.left_indent = Cm(0.74)

    doc.add_heading('3.3 Nested List', level=2)
    p = doc.add_paragraph()
    run = p.add_run('• Level 1 item')
    set_run_font(run)
    p.paragraph_format.left_indent = Cm(0.74)

    p = doc.add_paragraph()
    run = p.add_run('• Level 2 item')
    set_run_font(run)
    p.paragraph_format.left_indent = Cm(1.48)

    p = doc.add_paragraph()
    run = p.add_run('• Level 3 item')
    set_run_font(run)
    p.paragraph_format.left_indent = Cm(2.22)

    doc.add_heading('4. Table Styles', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Table Grid, centered, header background #D9D9D9, header centered, data left-aligned.')
    set_run_font(run, size=10, italic=True)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Column 1', 'Column 2', 'Column 3']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
        set_cell_shading(cell, 'D9D9D9')

    data = [
        ['Value 1', 'Value 2', 'Value 3'],
        ['Value 4', 'Value 5', 'Value 6'],
        ['Value 7', 'Value 8', 'Value 9'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text

    doc.add_heading('5. Code Block Styles', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Code block: Consolas 9pt, background #F5F5F5, left indent 0.5cm')
    set_run_font(run, size=10, italic=True)

    p = doc.add_paragraph()
    run = p.add_run('[python]')
    set_run_font(run, font_en='Consolas', size=9, italic=True)

    code_lines = [
        'def hello_world():',
        '    print("Hello, World!")',
        '    return True',
    ]
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, font_cn='Consolas', font_en='Consolas', size=9)
        p.paragraph_format.left_indent = Cm(0.5)
        add_shading(p, 'F5F5F5')

    doc.add_heading('6. Inline Code Styles', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Inline code: ')
    set_run_font(run)
    run = p.add_run('print("Hello")')
    set_run_font(run, font_cn='Consolas', font_en='Consolas', size=12)

    p = doc.add_paragraph()
    run = p.add_run('Inline code: Consolas 12pt, background #F0F0F0')
    set_run_font(run, size=10, italic=True)

    doc.add_heading('7. Blockquote Styles', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Blockquote: left border #6366F1, 1.5pt width, 1cm indent, italic')
    set_run_font(run, size=10, italic=True)

    p = doc.add_paragraph()
    run = p.add_run('This is quoted text in italic with a purple left border.')
    set_run_font(run, italic=True)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)

    doc.add_heading('8. Horizontal Rule Styles', level=1)

    p = doc.add_paragraph()
    run = p.add_run('Horizontal rule: bottom border #CCCCCC, 6pt spacing before/after')
    set_run_font(run, size=10, italic=True)

    p = doc.add_paragraph()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph('Content above the rule')
    doc.add_paragraph('Content below the rule')

    doc.add_heading('9. Font Specification Summary', level=1)

    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Element', 'CJK font', 'Latin font', 'Size']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
        set_cell_shading(cell, 'D9D9D9')

    data = [
        ['Body', 'SimSun', 'Times New Roman', '12pt'],
        ['Heading 1', 'SimSun', 'Times New Roman', '22pt'],
        ['Heading 2', 'SimSun', 'Times New Roman', '16pt'],
        ['Heading 3', 'SimSun', 'Times New Roman', '15pt'],
        ['Heading 4', 'SimSun', 'Times New Roman', '14pt'],
        ['Code block', 'Consolas', 'Consolas', '9pt'],
        ['Inline code', 'Consolas', 'Consolas', '12pt'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text

    doc.add_heading('10. Paragraph Specification Summary', level=1)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Property', 'Value', 'Notes']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
        set_cell_shading(cell, 'D9D9D9')

    data = [
        ['First-line indent', '0.74cm', '~two CJK character widths'],
        ['Line spacing', '1.5×', 'Improved readability'],
        ['Space before', '0pt', 'Compact layout'],
        ['Space after', '0pt', 'Compact layout'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text

    doc.add_heading('11. Page Break Rules', level=1)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Rule', 'Description']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
        set_cell_shading(cell, 'D9D9D9')

    data = [
        ['Page break before level-1 headings', 'Each level-1 heading starts a new page'],
        ['Other headings', 'Level 2 and below stay on the same page'],
        ['After cover page', 'Automatic page break after the cover'],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = cell_text

    doc.save(OUTPUT_PATH)
    print(f'Preview document created: {OUTPUT_PATH}')

if __name__ == '__main__':
    create_preview()
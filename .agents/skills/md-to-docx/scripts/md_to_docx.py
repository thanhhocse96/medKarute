#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import os
import sys
import shutil
from docx import Document
from docx.shared import Pt, Inches, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import Optional, List, Dict, Any, Tuple
from io import BytesIO

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    from markdown_normalizer import MarkdownNormalizer
    NORMALIZER_AVAILABLE = True
except ImportError:
    NORMALIZER_AVAILABLE = False

try:
    from version_manager import get_versioned_output_paths
    VERSION_MANAGER_AVAILABLE = True
except ImportError:
    VERSION_MANAGER_AVAILABLE = False

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SKILL_DIR = os.path.dirname(SCRIPT_DIR)
ASSETS_DIR = os.path.join(SKILL_DIR, 'assets')

if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

FRONT_MATTER_KEYS = ('title', 'date', 'version', 'audience')
CJK_FONT = 'SimSun'
DEFAULT_DATE_FMT = '%Y-%m-%d'


def extract_front_matter(content: str) -> Tuple[Dict[str, str], str]:
    """Parse YAML-like front matter and return metadata plus body."""
    meta: Dict[str, str] = {}
    match = re.match(r'^---\r?\n([\s\S]*?)\r?\n---\r?\n*', content)
    if not match:
        return meta, content

    fm = match.group(1)
    for key in FRONT_MATTER_KEYS:
        field = re.search(rf'^{re.escape(key)}:\s*(.+)$', fm, re.M)
        if field:
            meta[key] = field.group(1).strip().strip('"').strip("'")

    return meta, content[match.end():]


def split_title_subtitle(title: str) -> Tuple[str, str]:
    """Split title on em/en dash into main title and subtitle."""
    parts = re.split(r'\s*[—–]\s*', title, maxsplit=1)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    return title.strip(), ''


def titles_match(left: str, right: str) -> bool:
    """Loose match for skipping duplicate document titles."""
    left = left.strip()
    right = right.strip()
    if not left or not right:
        return False
    return left == right or left in right or right in left


def get_default_template_path() -> str:
    """Return absolute path to the default template in assets/."""
    return os.path.join(ASSETS_DIR, 'template.docx')

def resolve_template_path(template_path: str) -> str:
    """Resolve template path with relative paths and default asset fallback."""
    if template_path is None:
        return get_default_template_path()
    
    if os.path.isabs(template_path):
        return template_path
    
    if os.path.exists(template_path):
        return os.path.abspath(template_path)
    
    default_template = get_default_template_path()
    if os.path.exists(default_template):
        return default_template
    
    return template_path

class MarkdownParser:
    def __init__(self, enable_code_blocks: bool = True, md_file_path: str = None):
        self.elements = []
        self.current_table = None
        self.in_code_block = False
        self.code_content = []
        self.code_language = ''
        self.code_block_indent = ''
        self.skip_empty = False
        self.enable_code_blocks = enable_code_blocks
        self.md_file_path = md_file_path
        self.md_file_dir = os.path.dirname(md_file_path) if md_file_path else None
    
    def _parse_image(self, line: str) -> List[Dict[str, Any]]:
        """Parse a line containing images and return element dicts."""
        elements = []
        img_pattern = r'!\[([^\]]*)\]\(([^)\s]+)(?:\s+"([^"]*)")?\)'
        last_end = 0
        
        for match in re.finditer(img_pattern, line):
            if match.start() > last_end:
                text_before = line[last_end:match.start()].strip()
                if text_before:
                    elements.append({'type': 'paragraph', 'text': text_before})
            
            alt_text = match.group(1)
            img_path = match.group(2)
            img_title = match.group(3) if match.group(3) else ''
            
            if self.md_file_dir and not os.path.isabs(img_path) and not img_path.startswith('http'):
                img_path = os.path.join(self.md_file_dir, img_path)
            
            elements.append({
                'type': 'image',
                'alt': alt_text,
                'path': img_path,
                'title': img_title
            })
            last_end = match.end()
        
        if last_end < len(line):
            text_after = line[last_end:].strip()
            if text_after:
                elements.append({'type': 'paragraph', 'text': text_after})
        
        return elements if elements else [{'type': 'paragraph', 'text': line}]
    
    def parse(self, content: str) -> List[Dict[str, Any]]:
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            stripped_line = line.lstrip()
            if self.enable_code_blocks and stripped_line.startswith('```'):
                if self.in_code_block:
                    self.elements.append({
                        'type': 'code_block',
                        'content': '\n'.join(self.code_content),
                        'language': self.code_language
                    })
                    self.code_content = []
                    self.code_language = ''
                    self.code_block_indent = ''
                    self.in_code_block = False
                else:
                    self.in_code_block = True
                    self.code_language = stripped_line[3:].strip()
                    self.code_block_indent = line[:len(line) - len(stripped_line)]
                i += 1
                self.skip_empty = False
                continue
            
            if self.in_code_block:
                if self.code_block_indent and line.startswith(self.code_block_indent):
                    self.code_content.append(line[len(self.code_block_indent):])
                elif line.strip() == '':
                    self.code_content.append('')
                else:
                    self.code_content.append(line)
                i += 1
                self.skip_empty = False
                continue
            
            if self.current_table is not None:
                if line.strip().startswith('|'):
                    if '---' not in line:
                        self.current_table['rows'].append(line)
                    i += 1
                    continue
                else:
                    self._flush_table()
                    self.skip_empty = False
            
            if line.startswith('# '):
                self.elements.append({'type': 'title', 'text': line[2:].strip()})
                self.skip_empty = True
            elif line.startswith('## '):
                self.elements.append({'type': 'heading1', 'text': line[3:].strip()})
                self.skip_empty = True
            elif line.startswith('### '):
                self.elements.append({'type': 'heading2', 'text': line[4:].strip()})
                self.skip_empty = True
            elif line.startswith('#### '):
                self.elements.append({'type': 'heading3', 'text': line[5:].strip()})
                self.skip_empty = True
            elif line.startswith('##### '):
                self.elements.append({'type': 'heading4', 'text': line[6:].strip()})
                self.skip_empty = True
            elif line.startswith('###### '):
                self.elements.append({'type': 'heading5', 'text': line[7:].strip()})
                self.skip_empty = True
            elif re.match(r'^(\s*)[-*]\s', line):
                match = re.match(r'^(\s*)[-*]\s(.*)$', line)
                if match:
                    indent = match.group(1)
                    level = len(indent) // 2 + 1
                    text = match.group(2).strip()
                    self.elements.append({'type': 'bullet', 'text': text, 'level': level})
                self.skip_empty = False
            elif re.match(r'^(\s*)\d+\.\s', line):
                match = re.match(r'^(\s*)(\d+)\.\s(.*)$', line)
                if match:
                    indent = match.group(1)
                    level = len(indent) // 2 + 1
                    number = int(match.group(2))
                    text = match.group(3).strip()
                    self.elements.append({'type': 'ordered_item', 'text': text, 'number': number, 'level': level})
                self.skip_empty = False
            elif line.strip().startswith('|'):
                self.current_table = {'rows': [line]}
                self.skip_empty = False
            elif line.strip().startswith('>'):
                text = line.strip().lstrip('>').strip()
                self.elements.append({'type': 'blockquote', 'text': text})
                self.skip_empty = False
            elif line.strip() == '---' or line.strip() == '***' or line.strip() == '___':
                self.elements.append({'type': 'horizontal_rule'})
                self.skip_empty = False
            elif line.strip() == '':
                if not self.skip_empty:
                    self.elements.append({'type': 'empty'})
            else:
                if '![' in line and '](' in line:
                    img_elements = self._parse_image(line)
                    self.elements.extend(img_elements)
                else:
                    self.elements.append({'type': 'paragraph', 'text': line})
                self.skip_empty = False
            
            i += 1
        
        if self.current_table:
            self._flush_table()
        
        return self.elements
    
    def _parse_cell_content(self, cell_text: str) -> List[Dict[str, Any]]:
        """Parse table cell content that may mix text and images."""
        content = []
        img_pattern = r'!\[([^\]]*)\]\(([^)\s]+)(?:\s+"([^"]*)")?\)'
        last_end = 0
        
        for match in re.finditer(img_pattern, cell_text):
            if match.start() > last_end:
                text_before = cell_text[last_end:match.start()].strip()
                if text_before:
                    content.append({'type': 'text', 'value': text_before})
            
            alt_text = match.group(1)
            img_path = match.group(2)
            img_title = match.group(3) if match.group(3) else ''
            
            if self.md_file_dir and not os.path.isabs(img_path) and not img_path.startswith('http'):
                img_path = os.path.join(self.md_file_dir, img_path)
            
            content.append({
                'type': 'image',
                'alt': alt_text,
                'path': img_path,
                'title': img_title
            })
            last_end = match.end()
        
        if last_end < len(cell_text):
            text_after = cell_text[last_end:].strip()
            if text_after:
                content.append({'type': 'text', 'value': text_after})
        
        return content if content else [{'type': 'text', 'value': cell_text}]
    
    def _flush_table(self):
        if self.current_table and self.current_table['rows']:
            rows = self.current_table['rows']
            headers = [self._parse_cell_content(c.strip()) for c in rows[0].split('|') if c.strip()]
            data = []
            for row in rows[1:]:
                cells = [self._parse_cell_content(c.strip()) for c in row.split('|') if c.strip()]
                if cells:
                    data.append(cells)
            self.elements.append({'type': 'table', 'headers': headers, 'data': data})
        self.current_table = None


class TextFormatter:
    @staticmethod
    def parse_inline(text: str) -> List[Dict[str, Any]]:
        parts = []
        i = 0
        
        while i < len(text):
            if text[i] == '`' and (i == 0 or text[i-1] != '\\'):
                end_idx = text.find('`', i + 1)
                if end_idx != -1:
                    parts.append({'type': 'code', 'content': text[i+1:end_idx]})
                    i = end_idx + 1
                else:
                    parts.append({'type': 'text', 'content': text[i]})
                    i += 1
            elif text[i] == '[':
                link_match = re.match(r'\[([^\]]+)\]\(([^)\s]+)(?:\s+"([^"]*)")?\)', text[i:])
                if link_match:
                    link_text = link_match.group(1)
                    link_url = link_match.group(2)
                    link_title = link_match.group(3) if link_match.group(3) else ''
                    inner_parts = TextFormatter.parse_inline(link_text)
                    parts.append({'type': 'link', 'content': inner_parts, 'url': link_url, 'title': link_title})
                    i += len(link_match.group(0))
                else:
                    parts.append({'type': 'text', 'content': text[i]})
                    i += 1
            elif text[i:i+3] == '***':
                end_idx = text.find('***', i + 3)
                if end_idx != -1:
                    inner_content = text[i+3:end_idx]
                    inner_parts = TextFormatter.parse_inline(inner_content)
                    parts.append({'type': 'bold_italic', 'content': inner_parts})
                    i = end_idx + 3
                else:
                    parts.append({'type': 'text', 'content': text[i:i+3]})
                    i += 3
            elif text[i:i+2] == '~~':
                end_idx = text.find('~~', i + 2)
                if end_idx != -1:
                    inner_content = text[i+2:end_idx]
                    inner_parts = TextFormatter.parse_inline(inner_content)
                    parts.append({'type': 'strikethrough', 'content': inner_parts})
                    i = end_idx + 2
                else:
                    parts.append({'type': 'text', 'content': text[i:i+2]})
                    i += 2
            elif text[i:i+2] == '**':
                end_idx = text.find('**', i + 2)
                if end_idx != -1:
                    inner_content = text[i+2:end_idx]
                    inner_parts = TextFormatter.parse_inline(inner_content)
                    parts.append({'type': 'bold', 'content': inner_parts})
                    i = end_idx + 2
                else:
                    parts.append({'type': 'text', 'content': text[i:i+2]})
                    i += 2
            elif text[i] == '*' and (i == 0 or text[i-1] != '*'):
                end_idx = text.find('*', i + 1)
                if end_idx != -1 and text[end_idx-1] != '*':
                    inner_content = text[i+1:end_idx]
                    inner_parts = TextFormatter.parse_inline(inner_content)
                    parts.append({'type': 'italic', 'content': inner_parts})
                    i = end_idx + 1
                else:
                    parts.append({'type': 'text', 'content': text[i]})
                    i += 1
            elif text[i:i+4] == '<br>':
                parts.append({'type': 'br'})
                i += 4
            else:
                next_code = text.find('`', i)
                next_link = text.find('[', i)
                next_bold_italic = text.find('***', i)
                next_strikethrough = text.find('~~', i)
                next_bold = text.find('**', i)
                next_italic = text.find('*', i)
                next_br = text.find('<br>', i)
                
                candidates = []
                if next_code != -1:
                    candidates.append(('code', next_code))
                if next_link != -1:
                    candidates.append(('link', next_link))
                if next_bold_italic != -1:
                    candidates.append(('bold_italic', next_bold_italic))
                if next_strikethrough != -1:
                    candidates.append(('strikethrough', next_strikethrough))
                if next_bold != -1:
                    candidates.append(('bold', next_bold))
                if next_italic != -1:
                    candidates.append(('italic', next_italic))
                if next_br != -1:
                    candidates.append(('br', next_br))
                
                if not candidates:
                    parts.append({'type': 'text', 'content': text[i:]})
                    break
                
                candidates.sort(key=lambda x: x[1])
                next_type, next_pos = candidates[0]
                
                if next_pos > i:
                    parts.append({'type': 'text', 'content': text[i:next_pos]})
                i = next_pos
        
        return parts

    @staticmethod
    def parse_link(text: str) -> Optional[Dict[str, Any]]:
        match = re.match(r'\[([^\]]+)\]\(([^)\s]+)(?:\s+"([^"]*)")?\)', text)
        if match:
            return {
                'text': match.group(1),
                'url': match.group(2),
                'title': match.group(3) if match.group(3) else ''
            }
        return None


class DocxGenerator:
    def __init__(self, template_path: str):
        self.template_path = template_path
        self.doc = None
        self.first_heading1 = True
    
    def create_document(self, output_path: str):
        if self.template_path and os.path.exists(self.template_path):
            shutil.copy(self.template_path, output_path)
            self.doc = Document(output_path)
            for element in self.doc.element.body[:]:
                if element.tag.endswith('p') or element.tag.endswith('tbl'):
                    self.doc.element.body.remove(element)
        else:
            self.doc = Document()
            for section in self.doc.sections:
                section.page_height = Cm(29.7)
                section.page_width = Cm(21.0)
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
            self._setup_styles()
    
    def _setup_styles(self):
        styles_config = {
            'Heading 1': {'size': 22, 'bold': True, 'font_cn': CJK_FONT, 'font_en': 'Times New Roman'},
            'Heading 2': {'size': 16, 'bold': True, 'font_cn': CJK_FONT, 'font_en': 'Times New Roman'},
            'Heading 3': {'size': 15, 'bold': True, 'font_cn': CJK_FONT, 'font_en': 'Times New Roman'},
            'Heading 4': {'size': 14, 'bold': True, 'font_cn': CJK_FONT, 'font_en': 'Times New Roman'},
            'Heading 5': {'size': 14, 'bold': True, 'font_cn': CJK_FONT, 'font_en': 'Times New Roman'},
            'Heading 6': {'size': 12, 'bold': True, 'font_cn': CJK_FONT, 'font_en': 'Times New Roman'},
            'Normal': {'size': 12, 'bold': False, 'font_cn': CJK_FONT, 'font_en': 'Times New Roman'},
        }
        for style_name, config in styles_config.items():
            try:
                style = self.doc.styles[style_name]
                style.font.size = Pt(config['size'])
                style.font.bold = config['bold']
                style.font.name = config['font_en']
                style._element.rPr.rFonts.set(qn('w:eastAsia'), config['font_cn'])
            except KeyError:
                pass
    
    def _set_run_font(self, run, font_cn: str = CJK_FONT, font_en: str = 'Times New Roman', size: float = 12, bold: bool = False):
        run.font.size = Pt(size)
        run.font.name = font_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        run.bold = bold
    
    def _add_formatted_text(self, paragraph, text: str, base_size: float = 12):
        parts = TextFormatter.parse_inline(text)
        for part in parts:
            if part['type'] == 'text':
                run = paragraph.add_run(part['content'])
                self._set_run_font(run, size=base_size)
            elif part['type'] == 'bold':
                content = part['content']
                if isinstance(content, list):
                    for sub_part in content:
                        self._render_part(paragraph, sub_part, base_size, bold=True)
                else:
                    run = paragraph.add_run(content)
                    self._set_run_font(run, size=base_size, bold=True)
            elif part['type'] == 'italic':
                content = part['content']
                if isinstance(content, list):
                    for sub_part in content:
                        self._render_part(paragraph, sub_part, base_size, italic=True)
                else:
                    run = paragraph.add_run(content)
                    self._set_run_font(run, size=base_size)
                    run.italic = True
            elif part['type'] == 'bold_italic':
                content = part['content']
                if isinstance(content, list):
                    for sub_part in content:
                        self._render_part(paragraph, sub_part, base_size, bold=True, italic=True)
                else:
                    run = paragraph.add_run(content)
                    self._set_run_font(run, size=base_size, bold=True)
                    run.italic = True
            elif part['type'] == 'strikethrough':
                content = part['content']
                if isinstance(content, list):
                    for sub_part in content:
                        self._render_part(paragraph, sub_part, base_size, strike=True)
                else:
                    run = paragraph.add_run(content)
                    self._set_run_font(run, size=base_size)
                    run.font.strike = True
            elif part['type'] == 'code':
                run = paragraph.add_run(part['content'])
                run.font.size = Pt(base_size)
                run.font.name = 'Consolas'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F0F0F0')
                run._r.get_or_add_rPr().append(shading)
            elif part['type'] == 'link':
                link_url = part.get('url', '')
                link_content = part.get('content', [])
                link_text = self._extract_text_from_parts(link_content)
                if link_text:
                    self._add_hyperlink(paragraph, link_url, link_text, base_size)
            elif part['type'] == 'br':
                paragraph.add_run().add_break()
    
    def _render_part(self, paragraph, part: Dict[str, Any], base_size: float = 12, bold: bool = False, italic: bool = False, strike: bool = False):
        if part['type'] == 'text':
            run = paragraph.add_run(part['content'])
            self._set_run_font(run, size=base_size, bold=bold, italic=italic, strike=strike)
        elif part['type'] == 'bold':
            content = part['content']
            if isinstance(content, list):
                for sub_part in content:
                    self._render_part(paragraph, sub_part, base_size, bold=True, italic=italic, strike=strike)
            else:
                run = paragraph.add_run(content)
                self._set_run_font(run, size=base_size, bold=True)
                if italic:
                    run.italic = True
                if strike:
                    run.font.strike = True
        elif part['type'] == 'italic':
            content = part['content']
            if isinstance(content, list):
                for sub_part in content:
                    self._render_part(paragraph, sub_part, base_size, bold=bold, italic=True, strike=strike)
            else:
                run = paragraph.add_run(content)
                self._set_run_font(run, size=base_size)
                run.italic = True
                if bold:
                    run.bold = True
                if strike:
                    run.font.strike = True
        elif part['type'] == 'code':
            run = paragraph.add_run(part['content'])
            run.font.size = Pt(base_size)
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F0F0F0')
            run._r.get_or_add_rPr().append(shading)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        elif part['type'] == 'link':
            link_url = part.get('url', '')
            link_content = part.get('content', [])
            link_text = self._extract_text_from_parts(link_content)
            if link_text:
                self._add_hyperlink(paragraph, link_url, link_text, base_size)
    
    def _set_run_font(self, run, font_cn: str = CJK_FONT, font_en: str = 'Times New Roman', size: float = 12, bold: bool = False, italic: bool = False, strike: bool = False):
        run.font.size = Pt(size)
        run.font.name = font_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        run.bold = bold
        if italic:
            run.italic = True
        if strike:
            run.font.strike = True
    
    def _add_hyperlink(self, paragraph, url: str, text: str, base_size: float = 12):
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'SimSun')
        rPr.append(rFonts)
        
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(base_size * 2)))
        rPr.append(sz)
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    def _extract_text_from_parts(self, parts: List[Dict[str, Any]]) -> str:
        text_parts = []
        for part in parts:
            if part.get('type') == 'text':
                text_parts.append(part.get('content', ''))
            elif part.get('type') in ('bold', 'italic', 'bold_italic', 'strikethrough', 'code'):
                content = part.get('content')
                if isinstance(content, list):
                    text_parts.append(self._extract_text_from_parts(content))
                elif isinstance(content, str):
                    text_parts.append(content)
        return ''.join(text_parts)
    
    def add_cover_page(
        self,
        title: str,
        version: str = '',
        date: str = '',
        subtitle: str = '',
        audience: str = '',
    ):
        main_title, derived_subtitle = split_title_subtitle(title)
        if not subtitle:
            subtitle = derived_subtitle

        for _ in range(3):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(main_title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        if subtitle:
            p_sub = self.doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sub = p_sub.add_run(subtitle)
            run_sub.font.size = Pt(16)
            run_sub.font.name = 'SimSun'
            run_sub._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        for _ in range(14 if not subtitle else 10):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if version:
            run1 = p.add_run(f'Version: {version}')
            run1.font.size = Pt(12)
            run1.font.name = 'SimSun'
            run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            p.add_run('\n')

        if date:
            run2 = p.add_run(f'Date: {date}')
            run2.font.size = Pt(12)
            run2.font.name = 'SimSun'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        if audience:
            p.add_run('\n')
            run3 = p.add_run(f'Audience: {audience}')
            run3.font.size = Pt(12)
            run3.font.name = 'SimSun'
            run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        self.doc.add_paragraph()
        self.doc.add_page_break()

    def add_table_of_contents(self, entries: List[Tuple[int, str]]):
        """Insert a static table of contents from heading levels 1-3."""
        if not entries:
            return

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Table of Contents')
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        for level, text in entries:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74 * (level - 1))
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.bold = level <= 2
            run.font.size = Pt(14 if level == 1 else 12)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        self.doc.add_page_break()
    
    def add_title(self, text: str, is_first: bool = False, version: str = '', date: str = ''):
        if is_first:
            self.add_cover_page(text, version, date)
        else:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = 'SimSun'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    def add_heading(self, level: int, text: str):
        if level == 1:
            if not self.first_heading1:
                self.doc.add_page_break()
            self.first_heading1 = False
        
        p = self.doc.add_paragraph()
        
        # Disable page break before paragraph (template styles may enable it).
        p.paragraph_format.page_break_before = False

        try:
            style = self.doc.styles[f'Heading {level}']
            p.style = style
            # Re-apply after style assignment to override template defaults.
            p.paragraph_format.page_break_before = False
        except KeyError:
            pass
        
        pPr = p._p.get_or_add_pPr()
        for numPr in pPr.findall(qn('w:numPr')):
            pPr.remove(numPr)
        
        text = re.sub(r'^[\d\.]+\s*', '', text).strip()
        
        run = p.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        
        if level == 1:
            run.font.size = Pt(22)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(15)
        elif level == 4:
            run.font.size = Pt(14)
        elif level == 5:
            run.font.size = Pt(14)
    
    def add_paragraph(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        self._add_formatted_text(p, text)
    
    def add_paragraph_with_hanging_indent(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        self._add_formatted_text(p, text)
    
    def add_bullet(self, text: str, level: int = 1):
        p = self.doc.add_paragraph()
        indent_per_level = Cm(0.74)
        p.paragraph_format.left_indent = indent_per_level * level
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run('• ')
        self._set_run_font(run, size=12)
        self._add_formatted_text(p, text)
    
    def add_ordered_item(self, text: str, number: int, level: int = 1):
        p = self.doc.add_paragraph()
        indent_per_level = Cm(0.74)
        p.paragraph_format.left_indent = indent_per_level * level
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f'{number}. ')
        self._set_run_font(run, size=12)
        self._add_formatted_text(p, text)
    
    def add_table(self, headers: List[List[Dict[str, Any]]], data: List[List[List[Dict[str, Any]]]]):
        if not headers:
            return
        
        num_cols = len(headers)
        
        def get_cell_text_length(cell_content: List[Dict[str, Any]]) -> float:
            total_len = 0
            for item in cell_content:
                if item.get('type') == 'text':
                    text = item.get('value', '')
                    char_count = len(text)
                    chinese_count = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
                    total_len += char_count + chinese_count * 0.5
                elif item.get('type') == 'image':
                    total_len += 20
            return total_len
        
        col_widths = []
        for col_idx in range(num_cols):
            max_len = get_cell_text_length(headers[col_idx]) if col_idx < len(headers) else 0
            for row in data:
                if col_idx < len(row):
                    max_len = max(max_len, get_cell_text_length(row[col_idx]))
            col_widths.append(max(max_len, 5))
        
        total_width = sum(col_widths)
        page_width = 8464
        if total_width > 0:
            col_widths = [int(w * page_width / total_width) for w in col_widths]
        else:
            col_widths = [int(page_width / num_cols)] * num_cols
        
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for idx, width in enumerate(col_widths):
            for cell in table.columns[idx].cells:
                cell.width = Twips(width)
        
        def render_cell_content(cell, content: List[Dict[str, Any]], is_header: bool = False):
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            
            for item in content:
                if item.get('type') == 'text':
                    text = item.get('value', '')
                    if is_header:
                        self._add_formatted_text(p, text, base_size=10)
                    else:
                        self._add_formatted_text(p, text)
                elif item.get('type') == 'image':
                    img_path = item.get('path', '')
                    alt_text = item.get('alt', '')
                    
                    try:
                        if not PIL_AVAILABLE:
                            raise ImportError("PIL not available")
                        
                        if img_path.startswith('http://') or img_path.startswith('https://'):
                            if not REQUESTS_AVAILABLE:
                                raise ImportError("requests not available")
                            response = requests.get(img_path, timeout=10)
                            response.raise_for_status()
                            image_data = BytesIO(response.content)
                            img = Image.open(image_data)
                            image_stream = BytesIO()
                            img_format = img.format if img.format else 'PNG'
                            img.save(image_stream, format=img_format)
                            image_stream.seek(0)
                        else:
                            if not os.path.exists(img_path):
                                raise FileNotFoundError(f"Image file not found: {img_path}")
                            image_stream = img_path
                        
                        run = p.add_run()
                        run.add_picture(image_stream, width=Cm(10))
                    except Exception as e:
                        run = p.add_run(f"[Image: {alt_text}]")
                        run.font.size = Pt(9)
                        run.italic = True
        
        header_cells = table.rows[0].cells
        for i, header_content in enumerate(headers):
            render_cell_content(header_cells[i], header_content, is_header=True)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9D9D9')
            header_cells[i]._tc.get_or_add_tcPr().append(shading)
        
        for row_idx, row_data in enumerate(data):
            row = table.add_row()
            for i, cell_content in enumerate(row_data):
                if i < num_cols:
                    render_cell_content(row.cells[i], cell_content, is_header=False)
            if row_idx % 2 == 1:
                for cell in row.cells:
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'F2F7FB')
                    cell._tc.get_or_add_tcPr().append(shading)

        self.doc.add_paragraph()

    def add_code_block(self, code: str, language: str = ''):
        if language and language.lower() == 'mermaid':
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F5')
            p._p.get_or_add_pPr().append(shading)
            run = p.add_run('[Diagram: see source .md for interactive Mermaid diagram]')
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            run.italic = True
            return

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        p._p.get_or_add_pPr().append(shading)
        
        if language:
            lang_run = p.add_run(f'[{language}]\n')
            lang_run.font.size = Pt(9)
            lang_run.font.name = 'Consolas'
            lang_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            lang_run.italic = True
        
        code_run = p.add_run(code)
        code_run.font.size = Pt(9)
        code_run.font.name = 'Consolas'
        code_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    
    def add_blockquote(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        border = OxmlElement('w:pBdr')
        left_border = OxmlElement('w:left')
        left_border.set(qn('w:val'), 'single')
        left_border.set(qn('w:sz'), '12')
        left_border.set(qn('w:color'), '6366F1')
        border.append(left_border)
        p._p.get_or_add_pPr().append(border)
        
        self._add_formatted_text(p, text)
    
    def add_horizontal_rule(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)
    
    def add_image(self, img_path: str, alt_text: str = '', max_width: float = 15.0):
        """Add an image to the document.

        Args:
            img_path: Local path or URL
            alt_text: Alt text for caption / fallback
            max_width: Maximum width in centimeters
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        try:
            if not PIL_AVAILABLE:
                raise ImportError("PIL not available")
            
            if img_path.startswith('http://') or img_path.startswith('https://'):
                if not REQUESTS_AVAILABLE:
                    raise ImportError("requests not available")
                response = requests.get(img_path, timeout=10)
                response.raise_for_status()
                image_data = BytesIO(response.content)
                img = Image.open(image_data)
                image_stream = BytesIO()
                img_format = img.format if img.format else 'PNG'
                img.save(image_stream, format=img_format)
                image_stream.seek(0)
            else:
                if not os.path.exists(img_path):
                    raise FileNotFoundError(f"Image file not found: {img_path}")
                image_stream = img_path
            
            run = p.add_run()
            inline_shape = run.add_picture(image_stream, width=Cm(max_width))
            
            if alt_text:
                p_caption = self.doc.add_paragraph()
                p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_caption = p_caption.add_run(alt_text)
                run_caption.font.size = Pt(9)
                run_caption.font.name = 'SimSun'
                run_caption._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                run_caption.italic = True
                
        except Exception:
            error_msg = f"[Image not found: {img_path}]"
            run = p.add_run(error_msg)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.italic = True
            if alt_text:
                run_alt = p.add_run(f" (alt text: {alt_text})")
                run_alt.font.size = Pt(10)
                run_alt.italic = True
    
    def add_empty(self):
        self.doc.add_paragraph()
    
    def generate(
        self,
        elements: List[Dict[str, Any]],
        version: str = '',
        date: str = '',
        enable_code_blocks: bool = False,
        metadata: Optional[Dict[str, str]] = None,
        include_toc: bool = True,
    ):
        metadata = metadata or {}
        fm_title = metadata.get('title', '').strip()
        fm_version = metadata.get('version', '') or version
        fm_date = metadata.get('date', '') or date
        audience = metadata.get('audience', '')

        cover_from_front_matter = False
        skip_title_text = ''

        if fm_title:
            if not fm_date:
                fm_date = datetime.now().strftime(DEFAULT_DATE_FMT)
            self.add_cover_page(fm_title, fm_version, fm_date, audience=audience)
            cover_from_front_matter = True
            skip_title_text, _ = split_title_subtitle(fm_title)

        toc_entries: List[Tuple[int, str]] = []
        for elem in elements:
            if elem['type'] == 'heading1':
                toc_entries.append((1, elem['text']))
            elif elem['type'] == 'heading2':
                toc_entries.append((2, elem['text']))
            elif elem['type'] == 'heading3':
                toc_entries.append((3, elem['text']))

        if include_toc and toc_entries:
            self.add_table_of_contents(toc_entries)

        is_first_title = not cover_from_front_matter
        effective_version = fm_version or version
        effective_date = fm_date or date or datetime.now().strftime(DEFAULT_DATE_FMT)

        for elem in elements:
            if elem['type'] == 'title':
                if skip_title_text and titles_match(elem['text'], skip_title_text):
                    continue
                self.add_title(elem['text'], is_first_title, effective_version, effective_date)
                is_first_title = False

            elif elem['type'] == 'heading1':
                self.add_heading(1, elem['text'])
            
            elif elem['type'] == 'heading2':
                self.add_heading(2, elem['text'])
            
            elif elem['type'] == 'heading3':
                self.add_heading(3, elem['text'])
            
            elif elem['type'] == 'heading4':
                self.add_heading(4, elem['text'])
            
            elif elem['type'] == 'heading5':
                self.add_heading(5, elem['text'])
            
            elif elem['type'] == 'paragraph':
                if '<br>' in elem['text']:
                    self.add_paragraph_with_hanging_indent(elem['text'])
                else:
                    self.add_paragraph(elem['text'])
            
            elif elem['type'] == 'bullet':
                self.add_bullet(elem['text'], elem.get('level', 1))
            
            elif elem['type'] == 'ordered_item':
                self.add_ordered_item(elem['text'], elem['number'], elem.get('level', 1))
            
            elif elem['type'] == 'table':
                self.add_table(elem['headers'], elem['data'])
            
            elif elem['type'] == 'code_block' and enable_code_blocks:
                self.add_code_block(elem['content'], elem.get('language', ''))
            
            elif elem['type'] == 'blockquote':
                self.add_blockquote(elem['text'])
            
            elif elem['type'] == 'horizontal_rule':
                self.add_horizontal_rule()
            
            elif elem['type'] == 'image':
                self.add_image(elem['path'], elem.get('alt', ''))
            
            elif elem['type'] == 'empty':
                self.add_empty()
    
    def save(self, output_path: str):
        self.doc.save(output_path)


def convert_markdown_to_docx(
    markdown_content: str,
    output_path: str,
    template_path: str = None,
    version: str = '',
    date: str = '',
    normalize: bool = True,
    save_normalized: bool = True,
    normalized_output_path: Optional[str] = None,
    enable_code_blocks: bool = True,
    md_file_path: str = None,
    verbose: bool = False
) -> Tuple[str, List[str]]:
    fixes_applied = []
    
    resolved_template_path = resolve_template_path(template_path)
    if verbose:
        print(f"[Template] Resolved template path: {resolved_template_path}")
    
    metadata, markdown_body = extract_front_matter(markdown_content)
    if metadata.get('version') and not version:
        version = metadata['version']
    if metadata.get('date') and not date:
        date = metadata['date']

    if normalize and NORMALIZER_AVAILABLE:
        normalizer = MarkdownNormalizer(verbose=verbose)
        markdown_body = normalizer.normalize(markdown_body)
        fixes_applied = normalizer.get_fixes()

        if save_normalized and normalized_output_path:
            with open(normalized_output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_body)
            if verbose:
                print(f"[Normalizer] Saved normalized file to: {normalized_output_path}")

    parser = MarkdownParser(enable_code_blocks=enable_code_blocks, md_file_path=md_file_path)
    elements = parser.parse(markdown_body)

    generator = DocxGenerator(resolved_template_path)
    generator.create_document(output_path)
    generator.generate(
        elements,
        version,
        date,
        enable_code_blocks,
        metadata=metadata,
        include_toc=True,
    )
    generator.save(output_path)
    
    return output_path, fixes_applied


def convert_markdown_file(
    input_path: str,
    template_path: str = None,
    output_path: Optional[str] = None,
    version: str = '',
    date: str = '',
    normalize: bool = True,
    save_normalized: bool = True,
    enable_code_blocks: bool = True,
    verbose: bool = False,
    use_versioning: bool = True
) -> Tuple[str, str, List[str]]:
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    normalized_output_path = None
    actual_version = version
    
    if use_versioning and VERSION_MANAGER_AVAILABLE:
        version_info = get_versioned_output_paths(input_path, None, verbose)
        
        if output_path is None:
            output_path = version_info['docx_path']
        
        normalized_output_path = version_info['normalized_md_path']
        
        if not version:
            actual_version = f"V{version_info['version']}"
        
        if verbose:
            print(f"[VersionManager] Using version: {actual_version}")
    else:
        if output_path is None:
            output_path = os.path.splitext(input_path)[0] + '.docx'
        
        if normalize and save_normalized:
            base_name = os.path.splitext(input_path)[0]
            normalized_output_path = f"{base_name}_normalized.md"
    
    output_path, fixes_applied = convert_markdown_to_docx(
        content, output_path, template_path,
        actual_version, date, normalize, save_normalized,
        normalized_output_path, enable_code_blocks, input_path, verbose
    )
    
    return output_path, normalized_output_path, fixes_applied


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [template.docx] [output.docx] [version] [date]")
        print("Options:")
        print("  --no-normalize       Skip markdown normalization")
        print("  --no-save-norm       Don't save normalized markdown file")
        print("  --enable-code-blocks Enable code block rendering")
        print("  --no-versioning      Disable automatic version numbering")
        print("  --verbose            Show detailed processing info")
        print("")
        print("Note: If template.docx is not specified, the default template will be used.")
        sys.exit(1)
    
    args = sys.argv[1:]
    normalize = True
    save_normalized = True
    enable_code_blocks = True
    verbose = False
    use_versioning = True
    
    if '--no-normalize' in args:
        normalize = False
        args.remove('--no-normalize')
    
    if '--no-save-norm' in args:
        save_normalized = False
        args.remove('--no-save-norm')
    
    if '--enable-code-blocks' in args:
        enable_code_blocks = True
        args.remove('--enable-code-blocks')
    
    if '--verbose' in args:
        verbose = True
        args.remove('--verbose')
    
    if '--no-versioning' in args:
        use_versioning = False
        args.remove('--no-versioning')
    
    input_file = args[0] if len(args) > 0 else None
    template_file = args[1] if len(args) > 1 else None
    output_file = args[2] if len(args) > 2 else None
    version = args[3] if len(args) > 3 else ''
    date = args[4] if len(args) > 4 else ''
    
    if input_file:
        result, normalized_path, fixes = convert_markdown_file(
            input_file, template_file, output_file,
            version, date, normalize, save_normalized,
            enable_code_blocks, verbose, use_versioning
        )
        print(f"Document created: {result}")
        if normalized_path:
            print(f"Normalized MD: {normalized_path}")
        if fixes:
            print(f"Markdown fixes applied: {len(fixes)}")
